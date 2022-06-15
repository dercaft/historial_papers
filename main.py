import csv,re,os
import pickle
from docx import Document
import pickle

from translate import trans_para
from generate import EEBO_Controller,EEBO_Parser
from generate import OLL_Parser,OLL_Controller
TILES=["html"]
FILE_PATH="./papers2.csv"

if __name__=="__main__":
    lister=[]
    with open(FILE_PATH,"r", encoding="utf-8") as f:
        read=csv.reader(f)
        lister=[i for i in read]
    for index,i in enumerate(lister):
        # if index <10 : continue
        title,path=i[2],i[3]
        # if not "卡托的信" in title: continue
        if i[1]=="EEBO":
            parser=EEBO_Parser()
            controller=EEBO_Controller(path,parser)
        elif i[1]=="OLL":
            parser=OLL_Parser()
            controller=OLL_Controller(path,parser)
        else: continue
        controller.explore(controller.container)
        content=parser.article
        for i,t in enumerate(content):
            content[i]=re.sub("[ ]+"," ",t)
            content[i]=re.sub("(〈◊〉|〈◊◊〉)","",content[i])
            
        name=path.split("/")[-1]
        for t in TILES:
            if name.endswith(t):
                name=name.rstrip("."+t)
                break
    
        print("*"*20)
        print("We get: ",len(content))
        print(len(path),path)
        print(len(name),name)
        ## Txt format
        with open(os.path.join("./txt",name+".txt") ,"w+",encoding="utf-8") as f:
            for line in content:
                f.write(line)
                f.write("\n")
        ## Translation
        print("TITLE: ",title)
        zh_content=[]
        for i,para in enumerate(content):
            zh_para=trans_para(para) if len(para) else ""
            zh_content.append(zh_para)
            print(i,zh_para)
        name=path.split("/")[-1].split(".")[0]
        print("*"*20)
        ## pkl format
        filename=name+".pkl"
        with open("./ckpt/"+filename,"wb+") as f:
            pickle.dump(zh_content,f)
        ## Save to word
        en=content
        zh=zh_content
        n=name
        document=Document()
        para=document.add_heading(title)
        for j,(e,z) in enumerate(zip(en,zh)):
            document.add_paragraph(e)
            document.add_paragraph(z)
        document.save("result/{}.docx".format(n))
