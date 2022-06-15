import csv,re,os
import pickle
from docx import Document
import pickle

from translate import trans_para
from generate import EEBO_Controller,EEBO_Parser
from generate import OLL_Parser,OLL_Controller
TILES=["html"]
if __name__=="__main__":
    filename_="CARTO.txt"
    file_=open(os.path.join("./txt",filename_) ,"w+",encoding="utf-8")
    lister=[]
    with open("./papers.csv","r", encoding="utf-8") as f:
        read=csv.reader(f)
        lister=[i for i in read]
    for index,i in enumerate(lister):
        if index <10 : continue
        title,path=i[2],i[3]
        if not "卡托的信" in title: continue
        if i[1]=="EEBO":
            parser=EEBO_Parser()
            controller=EEBO_Controller(path,parser)
        elif i[1]=="OLL":
            parser=OLL_Parser()
            controller=OLL_Controller(path,parser)
        else: continue
        controller.explore(controller.container)
        content=parser.article
        for i,t in enumerate(content):
            content[i]=re.sub("[ ]+"," ",t)
            
        name=path.split("/")[-1]
        for t in TILES:
            if name.endswith(t):
                name=name.rstrip("."+t)
                break
    
        print("*"*20)
        print("We get: ",len(content))
        print(len(path),path)
        print(name)
        ## Txt format
        if not os.path.exists(os.path.join("./txt",name+".txt")):
            with open(os.path.join("./txt",name+".txt") ,"w+",encoding="utf-8") as f:
                for line in content:
                    f.write(line)
                    f.write("\n")
                    file_.write(line)
                    file_.write("\n")
        ## load checkpoint
        ## Translation
        zfilename="temp_"+name+".txt"
        ckpt_length=0
        if os.path.exists(zfilename):
            with open(zfilename,"r",encoding="utf-8") as f:
                ckpt_length=len(f.readlines())//2
        # temp file finished
        if ckpt_length>=len(content):continue

        zfile=open(zfilename,"a+",encoding="utf-8")
        print("TITLE: ",title)
        zh_content=[]
        filename=name+".pkl"
        if ckpt_length: # Load from checkpoint
            zh_content=pickle.load(open("./ckpt/"+filename,"rb"))

        for i,para in enumerate(content):
            if i<ckpt_length: continue
            try:
                zh_para=trans_para(para) if len(para) else ""
                zh_content.append(zh_para)
            except BaseException as e:
                print(repr(e))
                break
            else:
                zfile.write(para+"\n")
                zfile.write(zh_para+"\n")
                print(i,zh_para)
        print("*"*20)
        ## pkl format
        with open("./ckpt/"+filename,"wb+") as f:
            pickle.dump(zh_content,f)
        ## Save to word
        zh=zh_content
        if ckpt_length: # Load from checkpoint
            zh=pickle.load(open("./ckpt/"+filename,"rb"))
        en=content
        n=name
        document=Document()
        para=document.add_heading(title)
        for j,(e,z) in enumerate(zip(en,zh)):
            document.add_paragraph(e)
            document.add_paragraph(z)
        document.save("result/{}.docx".format(n))
