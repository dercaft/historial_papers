import os
from docx import Document
import generate
import pickle
import csv
path=os.path.dirname(os.path.abspath(__file__))
ckpt=os.path.join(path,"ckpt")
html=os.path.join(path,"files")
ckpt_files=os.listdir(ckpt)
names=[ i.split(".")[0] for i in ckpt_files]
ckpt_paths=[ os.path.join(ckpt,i) for i in ckpt_files]
html_files=[i for i in os.listdir(html) if i.split(".")[0] in names]
html_paths=[ os.path.join(html,i) for i in html_files]
print(len(ckpt_files))
print(len(html_files))
print([i for i in names if i+".html" not in html_files])
assert all([len(html_paths)==len(ckpt_paths), len(names)==len(html_paths)])

infos=[]
with open("./papers.csv","r", encoding="utf-8") as f:
    reader=csv.reader(f)
    infos=[i for i in reader]
info_T=[[],[],[],[]]
for inf in infos:
    info_T[0].append(inf[0])
    info_T[1].append(inf[1])
    info_T[2].append(inf[2])
    info_T[3].append(inf[3])
    
for i,(n,hp,cp) in enumerate(zip(names,html_paths,ckpt_paths)):
    en=generate.decodeEEBO(hp)
    zh=pickle.load(open(cp,'rb'))
    print("RAW: {}, Trans: {}".format(len(en),len(zh)))
    
    ii=info_T[3].index("files/"+n+".html")
    title=info_T[2][ii]
    
    document=Document()
    para=document.add_heading(title)
    for j,(e,z) in enumerate(zip(en,zh)):
        document.add_paragraph(e)
        document.add_paragraph(z)
    document.save("result/{}.docx".format(n))
    print("{}/{} Finished.".format(i+1,len(html_paths)))